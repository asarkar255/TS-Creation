from docx import Document
from docx.shared import Pt, RGBColor
import re

def add_heading(doc, text):
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(text)
    run.bold = True
    run.font.color.rgb = RGBColor(0, 0, 255)  # Blue
    run.underline = True
    run.font.size = Pt(14)
    paragraph.style = 'Normal'

def add_content(doc, text):
    doc.add_paragraph(text)

def extract_flowchart_text(ts_text: str) -> str:
    flowchart_section = re.search(r"12\. Flowchart\s*(.*?)\n(?=\d{1,2}\.|\Z)", ts_text, re.DOTALL)
    if flowchart_section:
        return flowchart_section.group(1).strip()
    return ""

def create_docx(ts_text: str, buffer):
    doc = Document()
    doc.add_heading('Technical Specification', level=1)

    lines = ts_text.splitlines()
    section_header_pattern = re.compile(r"^\s*(\d{1,2})\.\s*(.*?):\s*(.+)$")  # Matches "1. Title: Content"
    plain_header_pattern = re.compile(r"^\s*(\d{1,2})\.\s*(.+)$")            # Matches "1. Title"

    current_section = ""
    current_content = []

    for line in lines:
        line = line.strip()
        if not line:
            continue

        header_with_content = section_header_pattern.match(line)
        plain_header = plain_header_pattern.match(line)

        if header_with_content:
            # Write previous section if exists
            if current_section and current_content:
                add_heading(doc, current_section)
                add_content(doc, '\n'.join(current_content))
            # Split header and content
            current_section = f"{header_with_content.group(1)}. {header_with_content.group(2)}"
            current_content = [header_with_content.group(3)]
        elif plain_header:
            # Write previous section if exists
            if current_section and current_content:
                add_heading(doc, current_section)
                add_content(doc, '\n'.join(current_content))
            current_section = f"{plain_header.group(1)}. {plain_header.group(2)}"
            current_content = []
        else:
            current_content.append(line)

    # Final flush
    if current_section and current_content:
        add_heading(doc, current_section)
        add_content(doc, '\n'.join(current_content))

    doc.save(buffer)